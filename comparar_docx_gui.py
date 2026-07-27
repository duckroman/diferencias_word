#!/usr/bin/env python3
"""
comparar_docx_gui.py
---------------------
Interfaz gráfica para comparar hasta 7 documentos de Word (.docx) y generar
un reporte HTML con los párrafos que difieren entre ellos.

Estructura del reporte:
  • Columnas : Bloque | Pág.(aprox.) | <nombre doc 1> | <nombre doc 2> | …
  • Las palabras que cambian llevan fondo amarillo directamente sobre el texto,
    sin afectar las palabras que son idénticas.

REQUISITOS:
    pip install python-docx
    pip install tkinterdnd2      (opcional, habilita arrastrar y soltar)

EJECUTAR:
    python comparar_docx_gui.py
"""

import os
import sys
import html
import queue
import re
import threading
import subprocess
from pathlib import Path
from difflib import SequenceMatcher

import tkinter as tk
from tkinter import ttk, filedialog, messagebox

from docx import Document
from docx.oxml.ns import qn

try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    DND_DISPONIBLE = True
except ImportError:
    DND_DISPONIBLE = False

MAX_DOCS = 7
PALABRAS_POR_PAGINA_RESPALDO = 400

_MSG_LOG   = "LOG"
_MSG_OK    = "OK"
_MSG_ERROR = "ERROR"


# ==========================================================================
# LÓGICA DE COMPARACIÓN
# ==========================================================================

def _reparar_docx(ruta_docx):
    import zipfile, io
    RE_REL_NULL = re.compile(
        rb'<Relationship\b[^>]*\bTarget=["\'][^"\']*NULL[^"\']*["\'][^>]*/?>',
        re.IGNORECASE,
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(ruta_docx, 'r') as zin, \
         zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            nombre = item.filename
            if not nombre:
                continue
            partes = nombre.replace('\\', '/').split('/')
            if any(p.upper() == 'NULL' for p in partes):
                continue
            try:
                datos = zin.read(nombre)
            except Exception:
                continue
            if nombre.endswith('.rels'):
                datos = RE_REL_NULL.sub(b'', datos)
            zout.writestr(item, datos)
    buf.seek(0)
    return buf


def extraer_parrafos(ruta_docx):
    try:
        doc = Document(ruta_docx)
    except Exception as e:
        msg = str(e).lower()
        if "null" in msg or "no item named" in msg or "there is no item" in msg:
            doc = Document(_reparar_docx(ruta_docx))
        else:
            raise

    parrafos = []
    pagina_actual = 1
    tiene_marcas_reales = False
    num_parrafo = 0

    for p in doc.paragraphs:
        p_xml = p._p
        saltos  = p_xml.findall('.//' + qn('w:br') + "[@" + qn('w:type') + "='page']")
        renders = p_xml.findall('.//' + qn('w:lastRenderedPageBreak'))
        if renders:
            tiene_marcas_reales = True
            pagina_actual += len(renders)
        if saltos:
            tiene_marcas_reales = True
            pagina_actual += len(saltos)
        texto = p.text.strip()
        if texto:
            num_parrafo += 1
            parrafos.append({"num": num_parrafo, "pagina": pagina_actual, "texto": texto})

    if not tiene_marcas_reales:
        acumulado = 0
        for item in parrafos:
            acumulado += len(item["texto"].split())
            item["pagina"] = max(1, (acumulado - 1) // PALABRAS_POR_PAGINA_RESPALDO + 1)

    return parrafos, tiene_marcas_reales


def calcular_diferencias(lista_parrafos_por_doc):
    n = len(lista_parrafos_por_doc)
    textos_base = [item["texto"] for item in lista_parrafos_por_doc[0]]

    registros = []
    for doc_idx in range(1, n):
        textos_otro = [item["texto"] for item in lista_parrafos_por_doc[doc_idx]]
        sm = SequenceMatcher(None, textos_base, textos_otro, autojunk=False)
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag != "equal":
                registros.append((i1, i2, doc_idx, j1, j2))

    if not registros:
        return []

    rangos_base = sorted(set((r[0], r[1]) for r in registros))
    bloques_base = []
    cur_ini, cur_fin = rangos_base[0]
    for ini, fin in rangos_base[1:]:
        if ini <= cur_fin:
            cur_fin = max(cur_fin, fin)
        else:
            bloques_base.append((cur_ini, cur_fin))
            cur_ini, cur_fin = ini, fin
    bloques_base.append((cur_ini, cur_fin))

    resultado = []
    for b_ini, b_fin in bloques_base:
        filas = []
        if b_ini < b_fin:
            for k in range(b_ini, b_fin):
                filas.append((0, lista_parrafos_por_doc[0][k]))
        else:
            filas.append((0, None))

        for doc_idx in range(1, n):
            sub  = [r for r in registros if r[2] == doc_idx and r[0] < b_fin and r[1] > b_ini]
            sub += [r for r in registros if r[2] == doc_idx and r[0] == r[1] and b_ini <= r[0] <= b_fin]
            if not sub:
                continue
            j1 = min(r[3] for r in sub)
            j2 = max(r[4] for r in sub)
            if j1 < j2:
                for k in range(j1, j2):
                    filas.append((doc_idx, lista_parrafos_por_doc[doc_idx][k]))
            else:
                filas.append((doc_idx, None))

        resultado.append(filas)

    return resultado


# ==========================================================================
# DIFF DE PALABRAS
# ==========================================================================

def diff_palabras(texto_base, texto_otro):
    """
    Compara dos textos token a token (palabras + espacios).
    Retorna (tokens_base, tokens_otro): listas de (str, bool),
    donde bool=True indica que el token es diferente.
    """
    def tokenizar(t):
        return re.findall(r'\S+|\s+', t) if t else []

    pal_base = tokenizar(texto_base or "")
    pal_otro = tokenizar(texto_otro or "")
    sm = SequenceMatcher(None, pal_base, pal_otro, autojunk=False)

    tokens_base: list = []
    tokens_otro: list = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for t in pal_base[i1:i2]:
                tokens_base.append((t, False))
            for t in pal_otro[j1:j2]:
                tokens_otro.append((t, False))
        elif tag == "replace":
            for t in pal_base[i1:i2]:
                tokens_base.append((t, True))
            for t in pal_otro[j1:j2]:
                tokens_otro.append((t, True))
        elif tag == "delete":
            for t in pal_base[i1:i2]:
                tokens_base.append((t, True))
        elif tag == "insert":
            for t in pal_otro[j1:j2]:
                tokens_otro.append((t, True))

    return tokens_base, tokens_otro


def tokens_a_html(tokens):
    """
    Convierte lista de (token, es_diferente) en HTML.
    Los tokens iguales → texto plano.
    Los tokens diferentes → <span class="diff">texto</span>
    con fondo amarillo y negrita.
    """
    partes = []
    for texto, diferente in tokens:
        t_esc = html.escape(texto)
        if diferente:
            partes.append(f'<span class="diff">{t_esc}</span>')
        else:
            partes.append(t_esc)
    return "".join(partes)


# ==========================================================================
# CONSTRUCCIÓN DEL REPORTE HTML
# ==========================================================================

_CSS = """
/* ── Reset y base ─────────────────────────────────── */
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

body {
    font-family: 'Segoe UI', Calibri, Arial, sans-serif;
    font-size: 13px;
    background: #f0f4f8;
    color: #1a1a2e;
    padding: 24px 16px 48px;
}

h1 {
    font-size: 1.5rem;
    font-weight: 700;
    color: #1e3a5f;
    margin-bottom: 6px;
}

.subtitle {
    color: #4a6080;
    font-size: 0.9rem;
    margin-bottom: 20px;
}

/* ── Tabla ────────────────────────────────────────── */
.wrap {
    overflow-x: auto;
    border-radius: 10px;
    box-shadow: 0 4px 20px rgba(0,0,0,.12);
}

table {
    border-collapse: collapse;
    width: 100%;
    background: #fff;
    min-width: 700px;
}

/* Encabezado */
thead th {
    background: #2e5395;
    color: #fff;
    font-weight: 600;
    font-size: 0.82rem;
    text-transform: uppercase;
    letter-spacing: .04em;
    padding: 11px 10px;
    text-align: center;
    border: 1px solid #1d3c72;
    position: sticky;
    top: 0;
    z-index: 2;
    white-space: nowrap;
}
thead th:nth-child(1),
thead th:nth-child(2) {
    position: sticky;
    left: 0;
    z-index: 3;
}
thead th:nth-child(2) { left: 60px; }

/* Celdas */
tbody td {
    vertical-align: top;
    padding: 8px 10px;
    border: 1px solid #d0d8e4;
    line-height: 1.55;
    min-width: 280px;
    max-width: 420px;
    word-break: break-word;
}

/* Columnas fijas izquierda */
tbody td:nth-child(1) {
    position: sticky;
    left: 0;
    min-width: 50px;
    max-width: 60px;
    text-align: center;
    font-weight: 700;
    background: inherit;
    z-index: 1;
}
tbody td:nth-child(2) {
    position: sticky;
    left: 60px;
    min-width: 100px;
    max-width: 130px;
    text-align: center;
    font-size: 0.82rem;
    color: #4a6080;
    background: inherit;
    z-index: 1;
}

/* Filas alternadas por bloque */
tr.bloque-a { background: #eef4fb; }
tr.bloque-b { background: #f9f9f9; }

/* Fila separadora entre bloques */
tr.sep td {
    padding: 0;
    height: 5px;
    background: #c8d8ea;
    border: none;
}

/* ── Resaltado de diferencias ─────────────────────── */
span.diff {
    background: #ffeb3b;          /* amarillo */
    font-weight: 700;
    border-radius: 3px;
    padding: 0 1px;
}

/* Texto de párrafo ausente */
.ausente {
    color: #888;
    font-style: italic;
    font-size: 0.85rem;
}

/* ── Barra de info superior ──────────────────────── */
.info-bar {
    background: #fff;
    border: 1px solid #d0d8e4;
    border-radius: 8px;
    padding: 12px 16px;
    margin-bottom: 18px;
    font-size: 0.85rem;
    color: #333;
    line-height: 1.8;
}
.info-bar strong { color: #2e5395; }

/* ── Leyenda ─────────────────────────────────────── */
.leyenda {
    display: flex;
    align-items: center;
    gap: 8px;
    margin-bottom: 14px;
    font-size: 0.83rem;
    color: #555;
}
.leyenda .muestra {
    background: #ffeb3b;
    font-weight: 700;
    border-radius: 3px;
    padding: 1px 6px;
}
"""


def construir_reporte(rutas, bloques, marcas_reales_por_doc, ruta_salida):
    """
    Genera el reporte de diferencias como página HTML.

    Estructura pivot:
      Col 1 : Bloque
      Col 2 : Pág.(aprox.)
      Col 3…: Una columna por documento cargado (nombre del .docx como cabecera)

    Las palabras diferentes llevan <span class="diff"> con fondo amarillo.
    Las palabras iguales se muestran en texto plano.
    """
    n_docs  = len(rutas)
    nombres = [Path(r).name for r in rutas]

    # ── Encabezados de columna ──────────────────────────────────────────
    ths = "".join(f"<th>{html.escape(n)}</th>" for n in nombres)
    header_row = f"<tr><th>Bloque</th><th>Pág.(aprox.)</th>{ths}</tr>"

    # ── Info de documentos cargados ─────────────────────────────────────
    info_items = []
    for idx, ruta in enumerate(rutas):
        rol   = "BASE" if idx == 0 else f"Doc {idx + 1}"
        marca = "" if marcas_reales_por_doc[idx] else " <em>(paginación estimada)</em>"
        info_items.append(
            f"<strong>[{rol}]</strong> {html.escape(Path(ruta).name)}{marca}"
        )
    info_html = "<br>".join(info_items)

    # ── Cuerpo de la tabla ──────────────────────────────────────────────
    if not bloques:
        body_html = (
            f'<tr><td colspan="{2 + n_docs}" style="text-align:center;padding:20px;">'
            "No se encontraron diferencias entre los documentos.</td></tr>"
        )
    else:
        filas_html = []
        for num_bloque, filas in enumerate(bloques, start=1):
            clase = "bloque-a" if num_bloque % 2 == 1 else "bloque-b"

            # Agrupar ítems por doc_idx → lista de ítems
            items_por_doc: dict[int, list] = {}
            for doc_idx, item in filas:
                items_por_doc.setdefault(doc_idx, []).append(item)

            max_sub    = max((len(v) for v in items_por_doc.values()), default=1)
            base_items = items_por_doc.get(0, [])

            for sub in range(max_sub):
                # Col 1: Bloque
                bloque_cel = f'<td>{num_bloque}</td>' if sub == 0 else '<td></td>'

                # Col 2: Pág. del doc base en esta sub-fila
                pag_val = ""
                if sub < len(base_items) and base_items[sub] is not None:
                    it = base_items[sub]
                    pag_val = f"p.{it['pagina']} ¶{it['num']}"
                pag_cel = f"<td>{html.escape(pag_val)}</td>"

                # Texto de referencia para el diff
                texto_base_sub = ""
                if sub < len(base_items) and base_items[sub] is not None:
                    texto_base_sub = base_items[sub]["texto"]

                # Celdas de documentos
                celdas = []
                for doc_idx in range(n_docs):
                    doc_items = items_por_doc.get(doc_idx, [])

                    if sub >= len(doc_items):
                        celdas.append("<td></td>")
                        continue

                    item = doc_items[sub]

                    if item is None:
                        celdas.append('<td><span class="ausente">'
                                      "(párrafo ausente)</span></td>")
                        continue

                    texto_este = item["texto"]

                    if doc_idx == 0:
                        primer_otro = next(
                            (d for d in range(1, n_docs)
                             if sub < len(items_por_doc.get(d, []))
                             and items_por_doc[d][sub] is not None),
                            None,
                        )
                        if primer_otro is not None:
                            texto_ref = items_por_doc[primer_otro][sub]["texto"]
                            tokens, _ = diff_palabras(texto_este, texto_ref)
                        else:
                            tokens = [(texto_este, False)]
                    else:
                        _, tokens = diff_palabras(texto_base_sub, texto_este)

                    contenido = tokens_a_html(tokens)
                    celdas.append(f"<td>{contenido}</td>")

                celdas_html = "".join(celdas)
                filas_html.append(
                    f'<tr class="{clase}">{bloque_cel}{pag_cel}{celdas_html}</tr>'
                )

            # Fila separadora entre bloques
            filas_html.append(
                f'<tr class="sep"><td colspan="{2 + n_docs}"></td></tr>'
            )

        body_html = "\n".join(filas_html)

    # ── HTML completo ───────────────────────────────────────────────────
    pagina = f"""<!DOCTYPE html>
<html lang="es">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Comparación de documentos Word</title>
  <style>
{_CSS}
  </style>
</head>
<body>

<h1>Comparación de documentos Word</h1>
<p class="subtitle">
  Se compararon <strong>{n_docs}</strong> documentos usando
  <strong>{html.escape(nombres[0])}</strong> como documento base.
  Solo se muestran los párrafos que difieren.
</p>

<div class="info-bar">
  {info_html}
</div>

<div class="leyenda">
  <span>Leyenda:</span>
  <span class="muestra">palabra</span>
  <span>= texto que difiere entre documentos</span>
</div>

<div class="wrap">
  <table>
    <thead>
      {header_row}
    </thead>
    <tbody>
{body_html}
    </tbody>
  </table>
</div>

</body>
</html>
"""

    Path(ruta_salida).write_text(pagina, encoding="utf-8")


# ==========================================================================
# TRABAJO EN SEGUNDO PLANO
# ==========================================================================

def _comparar_en_hilo(archivos, salida, cola):
    try:
        listas, marcas = [], []
        for ruta in archivos:
            parrafos, tiene_marcas = extraer_parrafos(ruta)
            listas.append(parrafos)
            marcas.append(tiene_marcas)
            aviso = "" if tiene_marcas else "  [paginación estimada]"
            cola.put((_MSG_LOG, f"  ✓ {Path(ruta).name}: {len(parrafos)} párrafos{aviso}"))

        cola.put((_MSG_LOG, "Comparando párrafos…"))
        bloques = calcular_diferencias(listas)
        cola.put((_MSG_LOG, f"  Bloques de diferencia encontrados: {len(bloques)}"))

        cola.put((_MSG_LOG, "Generando reporte HTML…"))
        construir_reporte(archivos, bloques, marcas, salida)
        cola.put((_MSG_OK, salida))
    except Exception as exc:
        import traceback
        cola.put((_MSG_ERROR, traceback.format_exc()))


# ==========================================================================
# INTERFAZ GRÁFICA
# ==========================================================================

class App:
    POLL_MS = 100

    def __init__(self, root):
        self.root = root
        self.root.title("Comparador de documentos Word")
        self.root.geometry("740x580")
        self.root.minsize(620, 480)

        self.archivos = []
        self.ruta_salida = tk.StringVar(value=str(Path.cwd() / "comparacion.html"))
        self.ultima_ruta_generada = None
        self._cola = queue.Queue()

        self._construir_widgets()
        self._log("Listo. Agrega documentos .docx y presiona «Comparar y generar reporte».")

    def _construir_widgets(self):
        PAD = {"padx": 10, "pady": 6}

        ttk.Label(
            self.root,
            text="Comparador de documentos Word (.docx)",
            font=("Helvetica", 13, "bold"),
        ).pack(anchor="w", **PAD)

        ttk.Label(
            self.root,
            text="El primer documento de la lista es el BASE contra el que se comparan los demás.",
        ).pack(anchor="w", padx=10)

        zona = ttk.LabelFrame(self.root, text="Documentos (2 – 7)")
        zona.pack(fill="both", expand=True, **PAD)

        hint = ("Arrastra archivos .docx aquí  ·  o usa «Agregar archivos…»"
                if DND_DISPONIBLE else
                "Usa «Agregar archivos…» para cargar tus documentos .docx")
        self.zona_drop = tk.Label(
            zona, text=hint, relief="groove", bd=2,
            bg="#eef2f7", fg="#444", height=2, font=("Helvetica", 11),
        )
        self.zona_drop.pack(fill="x", padx=8, pady=(8, 4))

        if DND_DISPONIBLE:
            self.zona_drop.drop_target_register(DND_FILES)
            self.zona_drop.dnd_bind("<<Drop>>", self._al_soltar_archivos)
            self.zona_drop.bind("<Enter>", lambda _: self.zona_drop.config(bg="#d6e4f5"))
            self.zona_drop.bind("<Leave>", lambda _: self.zona_drop.config(bg="#eef2f7"))

        lista_frame = ttk.Frame(zona)
        lista_frame.pack(fill="both", expand=True, padx=8, pady=4)

        self.lista = tk.Listbox(lista_frame, selectmode="browse", height=7,
                                font=("Courier", 11))
        sb = ttk.Scrollbar(lista_frame, orient="vertical", command=self.lista.yview)
        self.lista.config(yscrollcommand=sb.set)
        self.lista.pack(side="left", fill="both", expand=True)
        sb.pack(side="left", fill="y")

        bf = ttk.Frame(zona)
        bf.pack(fill="x", padx=8, pady=(0, 8))
        for txt, cmd in [
            ("Agregar archivos…", self._agregar_archivos),
            ("Subir ↑",           lambda: self._mover(-1)),
            ("Bajar ↓",           lambda: self._mover(1)),
            ("Quitar",            self._quitar_seleccionado),
            ("Limpiar lista",     self._limpiar_lista),
        ]:
            ttk.Button(bf, text=txt, command=cmd).pack(side="left", padx=2)

        sf = ttk.LabelFrame(self.root, text="Archivo de salida")
        sf.pack(fill="x", **PAD)
        ttk.Entry(sf, textvariable=self.ruta_salida).pack(
            side="left", fill="x", expand=True, padx=(8, 4), pady=8)
        ttk.Button(sf, text="Examinar…", command=self._elegir_salida).pack(
            side="left", padx=(0, 8))

        af = ttk.Frame(self.root)
        af.pack(fill="x", **PAD)

        self.boton_comparar = ttk.Button(
            af, text="⚙  Comparar y generar reporte",
            command=self._ejecutar_comparacion)
        self.boton_comparar.pack(side="left", padx=(0, 6))

        self.boton_abrir = ttk.Button(
            af, text="🌐  Abrir reporte en navegador",
            command=self._abrir_reporte, state="disabled")
        self.boton_abrir.pack(side="left")

        ef = ttk.LabelFrame(self.root, text="Estado")
        ef.pack(fill="both", expand=True, **PAD)

        self.texto_estado = tk.Text(
            ef, height=6, state="disabled", wrap="word",
            font=("Courier", 10), bg="#f8f8f8")
        sb2 = ttk.Scrollbar(ef, orient="vertical", command=self.texto_estado.yview)
        self.texto_estado.config(yscrollcommand=sb2.set)
        self.texto_estado.pack(side="left", fill="both", expand=True, padx=(8, 0), pady=8)
        sb2.pack(side="left", fill="y", padx=(0, 8), pady=8)

        self.barra = ttk.Progressbar(self.root, mode="indeterminate")
        self.barra.pack(fill="x", padx=10, pady=(0, 10))

    def _log(self, msg):
        self.texto_estado.config(state="normal")
        self.texto_estado.insert("end", msg + "\n")
        self.texto_estado.see("end")
        self.texto_estado.config(state="disabled")

    def _refrescar_lista(self):
        self.lista.delete(0, "end")
        for i, ruta in enumerate(self.archivos, 1):
            etiq = "BASE" if i == 1 else f"#{i}  "
            self.lista.insert("end", f"[{etiq}]  {Path(ruta).name}")

    def _agregar_rutas(self, rutas):
        agregados = rechazados_ext = rechazados_cupo = 0
        for ruta in rutas:
            ruta = str(ruta).strip()
            if not ruta:
                continue
            if not ruta.lower().endswith(".docx"):
                rechazados_ext += 1
                continue
            if len(self.archivos) >= MAX_DOCS:
                rechazados_cupo += 1
                continue
            if ruta not in self.archivos:
                self.archivos.append(ruta)
                agregados += 1
        self._refrescar_lista()
        if agregados:
            self._log(f"Se agregaron {agregados} archivo(s). Total: {len(self.archivos)}.")
        if rechazados_ext:
            messagebox.showwarning("Archivos ignorados",
                f"{rechazados_ext} archivo(s) no son .docx y se ignoraron.")
        if rechazados_cupo:
            messagebox.showwarning("Máximo alcanzado",
                f"Solo se admiten {MAX_DOCS} documentos. "
                f"{rechazados_cupo} archivo(s) no se agregaron.")

    def _al_soltar_archivos(self, event):
        self._agregar_rutas(self.root.tk.splitlist(event.data))

    def _agregar_archivos(self):
        rutas = filedialog.askopenfilenames(
            title="Selecciona documentos .docx",
            filetypes=[("Documentos Word", "*.docx")])
        if rutas:
            self._agregar_rutas(rutas)

    def _quitar_seleccionado(self):
        sel = self.lista.curselection()
        if sel:
            del self.archivos[sel[0]]
            self._refrescar_lista()

    def _mover(self, delta):
        sel = self.lista.curselection()
        if not sel:
            return
        i = sel[0]
        j = i + delta
        if 0 <= j < len(self.archivos):
            self.archivos[i], self.archivos[j] = self.archivos[j], self.archivos[i]
            self._refrescar_lista()
            self.lista.selection_set(j)

    def _limpiar_lista(self):
        self.archivos.clear()
        self._refrescar_lista()

    def _elegir_salida(self):
        ruta = filedialog.asksaveasfilename(
            title="Guardar reporte como",
            defaultextension=".html",
            filetypes=[("Página HTML", "*.html"), ("Todos los archivos", "*.*")],
            initialfile="comparacion.html")
        if ruta:
            self.ruta_salida.set(ruta)

    def _ejecutar_comparacion(self):
        if len(self.archivos) < 2:
            messagebox.showerror("Faltan documentos",
                "Se necesitan al menos 2 documentos para comparar.")
            return
        salida = self.ruta_salida.get().strip()
        if not salida:
            messagebox.showerror("Sin ruta de salida",
                "Indica dónde guardar el reporte.")
            return

        self.boton_comparar.config(state="disabled")
        self.boton_abrir.config(state="disabled")
        self.barra.start(10)
        self._log("─" * 40)
        self._log(f"Iniciando comparación de {len(self.archivos)} documentos…")

        threading.Thread(
            target=_comparar_en_hilo,
            args=(list(self.archivos), salida, self._cola),
            daemon=True,
        ).start()
        self.root.after(self.POLL_MS, self._drenar_cola)

    def _drenar_cola(self):
        terminado = False
        try:
            while True:
                tipo, datos = self._cola.get_nowait()
                if tipo == _MSG_LOG:
                    self._log(datos)
                elif tipo == _MSG_OK:
                    self.ultima_ruta_generada = datos
                    self._log(f"✅ Reporte guardado en:\n   {datos}")
                    self._al_terminar_ok()
                    terminado = True
                elif tipo == _MSG_ERROR:
                    self._log(f"❌ ERROR:\n{datos}")
                    self._al_terminar_error(datos)
                    terminado = True
        except queue.Empty:
            pass
        if not terminado:
            self.root.after(self.POLL_MS, self._drenar_cola)

    def _al_terminar_ok(self):
        self.barra.stop()
        self.boton_comparar.config(state="normal")
        self.boton_abrir.config(state="normal")
        messagebox.showinfo("Listo", "El reporte HTML se generó correctamente.")

    def _al_terminar_error(self, msg):
        self.barra.stop()
        self.boton_comparar.config(state="normal")
        messagebox.showerror("Error al comparar", msg)

    def _abrir_reporte(self):
        if not self.ultima_ruta_generada:
            return
        try:
            if sys.platform == "darwin":
                subprocess.run(["open", self.ultima_ruta_generada], check=False)
            elif sys.platform.startswith("win"):
                os.startfile(self.ultima_ruta_generada)
            else:
                subprocess.run(["xdg-open", self.ultima_ruta_generada], check=False)
        except Exception as exc:
            messagebox.showerror("No se pudo abrir", str(exc))


# ==========================================================================
# PUNTO DE ENTRADA
# ==========================================================================

def main():
    if DND_DISPONIBLE:
        root = TkinterDnD.Tk()
    else:
        root = tk.Tk()
    App(root)
    root.mainloop()


if __name__ == "__main__":
    main()
